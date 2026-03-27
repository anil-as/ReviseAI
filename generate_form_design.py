import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def format_title(filename):
    # Remove extension and leading numbers like "01_"
    name = os.path.splitext(filename)[0]
    # Remove leading digits and underscores
    parts = name.split('_')
    if parts[0].isdigit():
        parts = parts[1:]
    
    # Capitalize and join
    title = ' '.join(p.capitalize() for p in parts)
    return title + " UI"

def create_form_design_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Form Design & UI Screenshots', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph("This document contains the form designs and UI screenshots for the ReviseAI application.\n")
    
    screenshots_dir = os.path.join(os.path.dirname(__file__), 'ui_screenshots')
    if not os.path.exists(screenshots_dir):
        print(f"Directory not found: {screenshots_dir}")
        return
        
    images = sorted([f for f in os.listdir(screenshots_dir) if f.endswith('.png') or f.endswith('.jpg')])
    
    for img in images:
        heading_text = format_title(img)
        h = doc.add_heading(heading_text, level=1)
        # Style heading
        if h.runs:
            h.runs[0].font.name = 'Times New Roman'
            h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
        img_path = os.path.join(screenshots_dir, img)
        try:
            doc.add_picture(img_path, width=Inches(6.0))
        except Exception as e:
            print(f"Error adding image {img}: {e}")
            doc.add_paragraph(f"[Image {img} could not be loaded]")
            
        doc.add_page_break()
        
    # the last page break might leave a blank page, but that's fine
    
    output_dir = os.path.join(os.path.dirname(__file__), 'other_docs')
    os.makedirs(output_dir, exist_ok=True)
    doc_path = os.path.join(output_dir, 'Form_Design_UI.docx')
    
    doc.save(doc_path)
    print(f"Document successfully created at: {doc_path}")

if __name__ == '__main__':
    create_form_design_doc()

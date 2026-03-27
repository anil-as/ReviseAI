import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Times New Roman'
    if level == 1:
        run.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_module_section(doc, title, intro, inputs, processing, outputs, responsibilities, inter_interfaces, ext_interfaces):
    add_heading(doc, title, level=2)
    
    p = doc.add_paragraph()
    p.add_run("Introduction: ").bold = True
    p.add_run(intro).font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run("Inputs: ").bold = True
    p.add_run(inputs).font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run("Processing: ").bold = True
    p.add_run(processing).font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run("Outputs: ").bold = True
    p.add_run(outputs).font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run("Functional Responsibilities: ").bold = True
    p.add_run(responsibilities).font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run("Interfaces with Other Modules: ").bold = True
    p.add_run(inter_interfaces).font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run("External Interfaces: ").bold = True
    p.add_run(ext_interfaces).font.name = 'Times New Roman'
    
    doc.add_paragraph()

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Detailed Modular Explanation', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    add_heading(doc, '1. Overview', level=1)
    add_paragraph(doc, 'This document provides a detailed modular explanation of the Smart Adaptive Revision Planning System (ReviseAI). It defines the core modules comprising the system, detailing their introductions, inputs, processing logic, outputs, functional responsibilities, and both internal and external interfaces.')
    
    add_heading(doc, '2. Modules', level=1)
    
    # Module 1
    add_module_section(doc, 
        title="2.1 Authentication and Authorization Module",
        intro="Responsible for handling user access securely. It manages user registration, login, token generation, and role-based access control.",
        inputs="User credentials (email, password), Registration details (name, role, etc.).",
        processing="Hashes passwords using algorithms like bcrypt, validates credentials against the database on login, and generates JWT (JSON Web Tokens) for managing active sessions. Enforces route protection based on user roles (Student, Instructor, Admin).",
        outputs="JWT access tokens, HTTP success/failure response codes and messages.",
        responsibilities="Securely authenticate users, manage session persistence, and prevent unauthorized access to protected API endpoints.",
        inter_interfaces="Interfaces heavily with the User Management Module to fetch and create user records, and is a dependency for all secure backend routers.",
        ext_interfaces="Interfaces with the frontend HTTP client (Axios) for token exchange."
    )
    
    # Module 2
    add_module_section(doc, 
        title="2.2 User and Enrollment Management Module",
        intro="Manages individual user profiles and handles the enrollment of students into specific subjects.",
        inputs="User profile updates (name, preferences), Subject IDs for enrollment requests.",
        processing="Creates links between user accounts and subjects in the database framework. Tracks enrollment dates and resolves many-to-many relationships to fetch comprehensive lists of subjects a specific student is actively enrolled in.",
        outputs="Detailed user profiles, lists of enrolled subjects, and updated user statistics.",
        responsibilities="Maintain up-to-date and accurate user personal information, handle student-subject access rights, and track enrollment lifecycles.",
        inter_interfaces="Authentication Module (for user ID resolution), Curriculum Module (to validate subject existence).",
        ext_interfaces="Frontend User Profile and Dashboard UI components."
    )
    
    # Module 3
    add_module_section(doc, 
        title="2.3 Curriculum (Subject & Topic) Module",
        intro="Manages the core hierarchical educational structure, encompassing Subjects and their underlying granular Topics.",
        inputs="Subject metadata (name, description, instructor ID), Topic metadata, and file uploads (PDFs, Documents) containing study context.",
        processing="Validates and stores educational structure in the database. Processes file uploads by saving them to the file system and extracting relevant text or references to form the context for the topic.",
        outputs="Hierarchical JSON responses containing Subjects and their nested Topics, file access paths/URLs.",
        responsibilities="Provide the educational backbone of the platform, logically organize study materials, and manage the storage mapping of uploaded documents.",
        inter_interfaces="User/Enrollment Module (to restrict access), Assessment Module, AI Generation Module (provides context for question generation).",
        ext_interfaces="Local File System or Cloud Storage for document persistence; Frontend Curriculum UI."
    )
    
    # Module 4
    add_module_section(doc, 
        title="2.4 Assessment and Spaced Repetition Module",
        intro="Handles the creation, retrieval, submission, grading, and scheduling of user assessments.",
        inputs="Topic IDs (to fetch relevant questions), User submitted answers (multiple choice selections, textual answers), current date.",
        processing="Validates submitted answers against correct answers. Calculates raw scores. Applies spaced repetition algorithms to determine and schedule the `next_revision_date` based on user performance. Stores assessment history to track progress.",
        outputs="Assessment results (score, accuracy), correct/incorrect indicators, updated revision schedules.",
        responsibilities="Administer valid tests, evaluate student understanding objectively, and implement the core spaced repetition logic to optimize long-term memory retention.",
        inter_interfaces="AI Generation Module (to retrieve generated questions), Curriculum Module (for linking to topics), User Module.",
        ext_interfaces="Frontend Assessment Interfaces and Assessment Flow UI."
    )
    
    # Module 5
    add_module_section(doc, 
        title="2.5 AI Question Generation & NLP Evaluator Module",
        intro="Automatically generates assessment questions from study materials and evaluates subjective/text-based answers using Natural Language Processing.",
        inputs="Topic textual context (extracted from uploaded documents), Student's descriptive textual answers.",
        processing="Constructs structured prompts encompassing the topic context and sends them to a Large Language Model (LLM) to generate various question types (MCQs, True/False, descriptive). Uses NLP embeddings and semantic similarity to evaluate subjective student answers against model answers.",
        outputs="Structured question sets in JSON format, calculated confidence or similarity scores for subjective grading.",
        responsibilities="Automate high-quality educational content creation dynamically, provide scalable and intelligent subjective grading capabilities.",
        inter_interfaces="Assessment Module (provides questions to be administered), Curriculum Module (retrieves context).",
        ext_interfaces="External LLM APIs (e.g., Google Generative AI / Gemini), NLP Python Libraries (SentenceTransformers)."
    )
    
    # Module 6
    add_module_section(doc, 
        title="2.6 Analytics and Dashboard Module",
        intro="Aggregates platform data to provide actionable insights, progress tracking, and statistical overviews for both students and instructors.",
        inputs="User IDs (Student or Instructor), Subject IDs, Date ranges.",
        processing="Performs complex database queries to aggregate assessment scores, completion rates, study streaks, and topic mastery levels. For instructors, aggregates overall class performance and active student counts.",
        outputs="Formatted statistical data (JSON) suitable for chart rendering.",
        responsibilities="Present actionable visual insights to users, track learning progress over time, and help instructors identify struggling students.",
        inter_interfaces="Assessment Module (for historical scores), User/Enrollment Module, Curriculum Module.",
        ext_interfaces="Frontend Dashboard UI (Charts, Graphs, and statistical cards)."
    )
    
    # Module 7
    add_module_section(doc, 
        title="2.7 Discussion / Subject Chat Module",
        intro="Provides a discussion board or messaging interface for each subject, allowing students and instructors to communicate, ask questions, and share information.",
        inputs="User textual messages, Subject ID.",
        processing="Stores messages in the database linked to the specific subject and the sender. Retrieves the message history for a subject, ensuring the user has the appropriate enrollment or instructor access rights to view or post.",
        outputs="A chronologically ordered list of chat messages for a given subject.",
        responsibilities="Facilitate communication between students and instructors within a subject, enable collaborative learning, and provide a forum for resolving doubts.",
        inter_interfaces="User Module (for sender details), Curriculum Module (to verify subject access and enrollment).",
        ext_interfaces="Frontend Chat/Discussion UI components."
    )
    
    # Module 8
    add_module_section(doc, 
        title="2.8 Client Interface (Frontend UI) Module",
        intro="The graphical user interface that users (students, instructors) interact with directly in their web browsers.",
        inputs="User interactions (clicks, form submissions, navigation), API responses from the backend.",
        processing="Renders responsive components, manages local application state (React Context/Redux), handles routing, and processes RESTful API calls to the backend using Axios. Handles client-side validation.",
        outputs="Rendered HTML/CSS/JS, API requests to the backend server.",
        responsibilities="Provide an intuitive, responsive, and seamless user experience. Act as the mediator between the human user and the backend REST APIs.",
        inter_interfaces="Connects logically to all backend API routers via the Service layer.",
        ext_interfaces="Browser APIs, Web Server hosting the static files."
    )

    doc_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'other_docs', 'Detailed_Modular_Explanation_Updated.docx')
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    doc.save(doc_path)
    print(f"Document saved to {doc_path}")

if __name__ == '__main__':
    create_doc()

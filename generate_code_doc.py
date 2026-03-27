import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def clean_code(code, lang):
    if lang == 'python':
        # Remove docstrings
        code = re.sub(r'\"\"\"[\s\S]*?\"\"\"', '', code)
        code = re.sub(r"\'\'\'[\s\S]*?\'\'\'", '', code)
        # Remove single line comments
        lines = []
        for line in code.split('\n'):
            line = re.sub(r'#.*', '', line)
            if line.strip():  # Skip empty lines
                lines.append(line.rstrip())
        return '\n'.join(lines)
    elif lang == 'javascript':
        # Remove multi-line comments
        code = re.sub(r'/\*.*?\*/', '', code, flags=re.DOTALL)
        # Remove single line comments (basic heuristic avoiding http://)
        lines = []
        for line in code.split('\n'):
            # simple exclusion for URLs
            if '//' in line and not 'http://' in line and not 'https://' in line:
                line = line.split('//')[0]
            if line.strip():
                lines.append(line.rstrip())
        return '\n'.join(lines)
    return code

def create_code_document():
    files_to_include = [
        # Backend
        ('backend/models.py', 'Database Models (models.py)', 'python'),
        ('backend/schemas.py', 'Pydantic Schemas (schemas.py)', 'python'),
        ('backend/routers/auth_router.py', 'Authentication Router (auth_router.py)', 'python'),
        ('backend/routers/user_router.py', 'User Router (user_router.py)', 'python'),
        ('backend/routers/subject_router.py', 'Subject Router (subject_router.py)', 'python'),
        ('backend/routers/topic_router.py', 'Topic Router (topic_router.py)', 'python'),
        ('backend/routers/assessment_router.py', 'Assessment Router (assessment_router.py)', 'python'),
        ('backend/services/question_generator.py', 'AI Question Generator Service (question_generator.py)', 'python'),
        ('backend/services/nlp_evaluator.py', 'NLP Evaluator Service (nlp_evaluator.py)', 'python'),
        ('backend/routers/dashboard_router.py', 'Dashboard API Router (dashboard_router.py)', 'python'),
        
        # Frontend
        ('frontend/src/App.js', 'Main App Entry (App.js)', 'javascript'),
        ('frontend/src/context/AuthContext.js', 'Authentication Context (AuthContext.js)', 'javascript'),
        ('frontend/src/context/ThemeContext.js', 'Theme Context (ThemeContext.js)', 'javascript'),
        ('frontend/src/pages/student/StudentDashboard.js', 'Student Dashboard Page (StudentDashboard.js)', 'javascript'),
        ('frontend/src/pages/student/StudentAssessment.js', 'Student Assessment Page (StudentAssessment.js)', 'javascript'),
        ('frontend/src/pages/instructor/InstructorDashboard.js', 'Instructor Dashboard Page (InstructorDashboard.js)', 'javascript'),
        ('frontend/src/pages/instructor/InstructorTopicDetails.js', 'Instructor Topic Details Page (InstructorTopicDetails.js)', 'javascript'),
        ('frontend/src/components/Navbar.js', 'Navigation Bar Component (Navbar.js)', 'javascript'),
        ('frontend/src/components/Sidebar.js', 'Sidebar Component (Sidebar.js)', 'javascript'),
        ('frontend/src/services/api.js', 'API Axios Client (api.js)', 'javascript'),
        ('frontend/src/services/assessmentService.js', 'Assessment Service (assessmentService.js)', 'javascript')
    ]
    
    root_dir = os.path.dirname(__file__)
    output_dir = os.path.join(root_dir, 'other_docs')
    os.makedirs(output_dir, exist_ok=True)
    doc_path = os.path.join(output_dir, 'Important_Source_Code.docx')
    
    doc = Document()
    
    title = doc.add_heading('Important Source Code Portions', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_page_break()
    
    for i, (rel_path, heading_text, lang) in enumerate(files_to_include, start=1):
        file_path = os.path.join(root_dir, rel_path)
        
        # Add heading
        h = doc.add_heading(f"{i}. {heading_text}", level=1)
        h.runs[0].font.name = 'Times New Roman'
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            stripped_content = clean_code(content, lang)
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(stripped_content)
            run.font.name = 'Courier New'
            run.font.size = Pt(8) # smaller font to fit well
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"Error: {rel_path} not found.")
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 0, 0)
            
        if i < len(files_to_include):
            doc.add_page_break()
            
    try:
        doc.save(doc_path)
        print(f"Document successfully saved to {doc_path}")
    except Exception as e:
        print(f"Failed to save document: {e}")

if __name__ == '__main__':
    create_code_document()
